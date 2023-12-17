import glob
import io
import json
import os
import random
from typing import List

from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm
from openpyxl import load_workbook
from tools.img_utils import base64_to_image_raw

VOUCHER_INFO = {
    "餐费": "5401001007998999",
    "业务招待费": "5401001007030"
}


def split_list(lst, n):
    """将列表lst拆分成n个元素的子列表"""
    return [lst[i:i + n] for i in range(0, len(lst), n)]


def makedir(path: str):
    if os.path.dirname(path) != "":
        os.makedirs(os.path.dirname(path), exist_ok=True)


def __json_read(json_path: str = "../data/json/*.json", sample_num: int = 20) -> list:
    """读取发票信息json文件"""
    invoice_list = []
    for json_path in glob.glob(json_path):
        with open(json_path, "r", encoding="utf-8") as f:
            invoice_dict = json.loads("".join(f.readlines()))
            invoice_list.append({
                "date": str(invoice_dict["invoice_info"]["date"]),
                "id": invoice_dict["invoice_info"]["id"],
                "code": invoice_dict["code"],
                "money": invoice_dict["invoice_info"]["total"],
                "img_b64": invoice_dict["verify_info"]["img_b64"],
            })
    return [random.choice(invoice_list) for _ in range(sample_num)]


def create_expense_ledger_xlsx(invoice_list: List[dict], claimant: str = "报销人", save_path: str = "费用报销台账.xlsx",
                               xlsx_template: str = "templates"):
    """
    创建费用报销台账

    :param invoice_list: 发票信息列表，列表内字典应包含以下字段
        'date': 发票日期    eg: '20201101'
        'id': 发票号码
        'money': 发票金额
    :param claimant: 报销人
    :param save_path: 费用报销台账存储路径
    :param xlsx_template: xlsx 模板文件夹路径
    """
    invoice_list = split_list(invoice_list, 12)
    for index, inv_list in enumerate(invoice_list):
        wb = load_workbook(filename=os.path.join(xlsx_template, '费用报销台账.xlsx'))
        sheet_ranges = wb['费用报销台账']
        for r_index, info in enumerate(inv_list):
            if r_index + 1 <= 12:
                sheet_ranges.cell(row=r_index + 4, column=1).value = r_index + 1  # 序号
                sheet_ranges.cell(row=r_index + 4,
                                  column=2).value = f"{info['date'][:4]}.{info['date'][4:6]}.{info['date'][6:]}"  # 日期
                # sheet_ranges.cell(row=r_index + 4, column=3).value = ""     # 事由
                sheet_ranges.cell(row=r_index + 4, column=4).value = "电子发票"  # 发票类型 []
                sheet_ranges.cell(row=r_index + 4, column=5).value = info['id']  # 发票号码(8位)
                sheet_ranges.cell(row=r_index + 4, column=6).value = info['money']  # 发票金额
                sheet_ranges.cell(row=r_index + 4, column=7).value = claimant  # 报销人
                # sheet_ranges.cell(row=r_index + 4, column=8).value = ""     # 备注
        # 写结果
        makedir(save_path)
        wb.save(save_path if len(
            invoice_list) == 1 else f"{os.path.splitext(save_path)[0]}_{index}{os.path.splitext(save_path)[1]}")


def create_invoice_verification_docx(invoice_list: List[dict], save_path: str = "发票真伪查询.docx"):
    """
    创建发票真伪查询docx (将selenium截取的查询结果png图片逐个置入docx内)

    :param invoice_list: 发票信息列表，列表内字典应包含以下字段
        'img_b64': 发票查验结果图片base64, png格式
    :param save_path: 发票真伪查询存储路径
    """

    doc = Document()
    # 窄边框
    doc.sections[0].top_margin = Cm(1.27)
    doc.sections[0].bottom_margin = Cm(1.27)
    doc.sections[0].left_margin = Cm(1.27)
    doc.sections[0].right_margin = Cm(1.27)
    # A4纸
    doc.sections[0].page_height = Cm(29.7)
    doc.sections[0].page_width = Cm(21)
    doc.sections[0].orientation = WD_ORIENTATION.LANDSCAPE  # 设置为纵向　　

    for info in invoice_list:
        img_stream = io.BytesIO(base64_to_image_raw(info['img_b64']))
        pic_add = doc.add_picture(img_stream, width=Cm(18.45))
        pic_add.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    makedir(save_path)
    doc.save(save_path)


def create_meal_expense_docx(invoice_list: List[dict], claimant: str = "报销人", save_path: str = "费用报销台账.xlsx",
                             docx_template: str = "templates"):
    """
    创建工作餐费报销明细表

    :param invoice_list: 发票信息列表，列表内字典应包含以下字段
        'date': 发票日期    eg: '20201101'
        'code': 发票代码
        'id': 发票号码
        'money': 发票金额
    :param claimant: 报销人
    :param save_path: 费用报销台账存储路径
    :param docx_template: docx 模板文件夹路径
    """
    invoice_list = split_list(invoice_list, 17)
    for index, inv_list in enumerate(invoice_list):
        # 读取模板文件
        doc = Document(os.path.join(docx_template, "工作餐费报销明细表.docx"))
        tabel = doc.tables[0]
        for r_index, info in enumerate(inv_list):
            if r_index + 1 <= 17:
                # 序号
                num_ = tabel.cell(row_idx=r_index + 1, col_idx=0)
                num_.text = str(r_index + 1)
                num_.paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                # 报销人姓名
                name_ = tabel.cell(row_idx=r_index + 1, col_idx=1)
                name_.text = str(claimant)
                name_.paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                # 发票代码
                code_ = tabel.cell(row_idx=r_index + 1, col_idx=2)
                code_.text = str(info['code'])
                code_.paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                # 发票号码(8位)
                id_ = tabel.cell(row_idx=r_index + 1, col_idx=3)
                id_.text = str(info['id'])
                id_.paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                # 日期
                date_ = tabel.cell(row_idx=r_index + 1, col_idx=4)
                date_.text = str(f"{info['date'][:4]}.{info['date'][4:6]}.{info['date'][6:]}")
                date_.paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                # 发票金额
                money_ = tabel.cell(row_idx=r_index + 1, col_idx=5)
                money_.text = str(info['money'])
                money_.paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                # 事由
                # reason_ = tabel.cell(row_idx=r_index + 1, col_idx=6)
                # reason_.text = str("")
                # reason_.paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                # 备注
                # remark_ = tabel.cell(row_idx=r_index + 1, col_idx=7)
                # remark_.text = str("")
                # remark_.paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        # 写结果
        makedir(save_path)
        doc.save(save_path if len(
            invoice_list) == 1 else f"{os.path.splitext(save_path)[0]}_{index}{os.path.splitext(save_path)[1]}")


def create_expense_voucher_xlsx(invoice_list: List[dict], claimant: str = "报销人",
                                save_path: str = "费用报销台账.xlsx",
                                xlsx_template: str = "templates"):
    """
    创建现金支出面单

    :param invoice_list: 发票信息列表，列表内字典应包含以下字段
        'money': 发票金额
    :param claimant: 报销人
    :param save_path: 现金支出面单存储路径
    :param xlsx_template: xlsx 模板文件夹路径
    """
    invoice_list = split_list(invoice_list, 8)
    for index, inv_list in enumerate(invoice_list):
        wb = load_workbook(filename=os.path.join(xlsx_template, '现金支出面单.xlsx'))
        sheet_ranges = wb['现金支出 （填写）']
        sheet_ranges.cell(row=10, column=4).value = claimant  # 报销人
        for inv_index, row in enumerate(range(21, 35, 2)):
            if inv_index + 1 <= len(inv_list):
                sheet_ranges.cell(row=row, column=3).value = "餐费"  # 摘要
                sheet_ranges.cell(row=row, column=7).value = inv_list[inv_index]["money"]  # 金额
                sheet_ranges.cell(row=row, column=9).value = "C"  # 借/贷 A/C
                sheet_ranges.cell(row=row, column=10).value = VOUCHER_INFO["餐费"]  # 科目代码
                sheet_ranges.cell(row=row, column=11).value = str(3650)  # Cost center
        # 写结果
        makedir(save_path)
        wb.save(save_path if len(
            invoice_list) == 1 else f"{os.path.splitext(save_path)[0]}_{index}{os.path.splitext(save_path)[1]}")


if __name__ == '__main__':
    invoice_lst = __json_read(json_path="../data/json/*.json", sample_num=3)
    # create expense ledger xlsx / 费用报销台账.xlsx
    create_expense_ledger_xlsx(invoice_list=invoice_lst, claimant="测试人", save_path="../export/费用报销台账.xlsx",
                               xlsx_template="../templates")

    # create invoice verification docx / 发票真伪查询.docx
    create_invoice_verification_docx(invoice_list=invoice_lst, save_path="../export/发票真伪查询.docx")

    # create meal expense docx / 工作餐费报销明细表.docx
    create_meal_expense_docx(invoice_list=invoice_lst, claimant="测试人", save_path="../export/工作餐费报销明细表.docx",
                             docx_template="../templates")

    # create expense voucher xlsx / 现金支出面单.xlsx
    create_expense_voucher_xlsx(invoice_list=invoice_lst, claimant="测试人", save_path="../export/现金支出面单.xlsx",
                                xlsx_template="../templates")
